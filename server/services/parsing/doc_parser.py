import io
import logging
from typing import Dict, Any
from docx import Document

logger = logging.getLogger(__name__)

def parse_docx_bytes(docx_bytes: bytes) -> Dict[str, Any]:
    """Парсит DOCX: возвращает текст по параграфам и таблицы."""
    doc = Document(io.BytesIO(docx_bytes))
    
    paragraphs = []
    tables = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)
        if table_data:
            tables.append(table_data)
    
    return {
        "pages_text": paragraphs,
        "tables": tables,
        "page_count": len(paragraphs),
        "has_tables": len(tables) > 0
    }
